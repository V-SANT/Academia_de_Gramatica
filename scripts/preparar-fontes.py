# -*- coding: utf-8 -*-
"""Prepara as fontes para transcrever unidades novas.

Os dois PDF do livro são digitalizações puras (uma imagem por página, sem
camada de texto) comprimidas em JBIG2 — o pypdf não as decodifica, o PDFium
sim. Este script transforma cada página numa imagem legível e extrai a chave
dos exercícios do .docx para texto simples.

    pip install pypdfium2 python-docx
    python scripts/preparar-fontes.py <pasta-com-as-fontes> <pasta-de-saida>

Espera encontrar na pasta das fontes:
    Gramática Ativa - 1-1-108.pdf          (as 108 páginas do livro)
    Gramatica_Ativa_1_Chave_Corrigida.docx (a chave dos exercícios)

Escreve em <pasta-de-saida>:
    paginas/pdf001.png … pdf108.png
    chave.txt
"""
import glob
import io
import os
import sys

import docx
import pypdfium2 as pdfium

fontes = sys.argv[1] if len(sys.argv) > 1 else "."
saida = sys.argv[2] if len(sys.argv) > 2 else "fontes"

os.makedirs(os.path.join(saida, "paginas"), exist_ok=True)


def encontrar(padrao):
    achados = glob.glob(os.path.join(fontes, padrao))
    if not achados:
        sys.exit(f"não encontrei nada com o padrão {padrao!r} em {fontes!r}")
    return achados[0]


# Páginas do livro -> PNG. escala 2.2 dá ~1300x1850, que se lê bem.
livro = encontrar("*1-1-108.pdf")
pdf = pdfium.PdfDocument(livro)
for i in range(len(pdf)):
    destino = os.path.join(saida, "paginas", f"pdf{i + 1:03d}.png")
    pdf[i].render(scale=2.2).to_pil().convert("L").save(destino, optimize=True)
print(f"{len(pdf)} páginas renderizadas em {saida}/paginas/")

# Chave dos exercícios -> texto
d = docx.Document(encontrar("*Chave*.docx"))
texto = "\n".join(p.text.strip() for p in d.paragraphs if p.text.strip())
io.open(os.path.join(saida, "chave.txt"), "w", encoding="utf-8", newline="\n").write(texto)
print(f"chave escrita em {saida}/chave.txt ({len(texto)} caracteres)")
